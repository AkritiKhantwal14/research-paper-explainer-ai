from docx import Document

def save_docx(content, filename):
    doc=Document()
    doc.add_heading("Research Paper Explainer", level=1)
    doc.add_paragraph(content)
    doc.save(filename)