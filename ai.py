import urllib.request#send requests to websites or api's
import json
from pdf_reader import extract_pdf_text
from chunker import chunk_text
from prompts import(simple_summary_prompt,fun_explanation,quiz, key_takeaways,beginner,literature_review_prompt, technical_prompt, research_gap_prompt,recommendation_prompt,citation_prompt,metadata_prompt)
from docx import Document
from report_generator import save_docx
from retriever import create_vector_store, retrieve_chunks

def generate_response(prompt):
    
    data = json.dumps({
        "model": "llama3",
        "prompt": prompt,
        "stream":False
    }).encode("utf-8")#convert to bytecode

    req= urllib.request.Request(
        "http://localhost:11434/api/generate",#send request to ollama running on my computer
        data=data,#json data
        headers={"Content-Type": "application/json"}#sending json data
    )
    print("Generating response")
    response=urllib.request.urlopen(req)#send request
    result=json.loads(response.read())
    return result["response"]

print("Reasearch Paper Explainer")

print("1.Simple Summary")
print("2.Fun Explanation")
print("3.Quiz Generation")
print("4.Key Takeaways")
print("5.Beginner Explanation")
print("6.Technical Explanation")
print("7.Literature Review")
print("8.Compare Research Papers")
print("9.Multi Paper Literature Review")
print("10.Research Gap Analysis")
print("11.Related Paper Recommendations")
print("12.Citation Generator")

choice = input("What would you like?")

if choice in["1","2","3","4","5","6","7","10","11","12","13"]:
    pdf_path=input("Enter PDF file path: ")
    article=extract_pdf_text(pdf_path)



if choice=="1": 
    prompt=simple_summary_prompt(article)


elif choice=="2": 
    prompt= fun_explanation(article)


elif choice=="3": 
    prompt= quiz(article)


elif choice=="4": 
    prompt= key_takeaways(article)


elif choice=="5": 
    prompt= beginner(article)


elif choice=="6": 
    prompt=technical_prompt(article)


elif choice=="7":
    prompt=literature_review_prompt(article)


elif choice=="8":
    
    num_papers=int(input("How many papers do you want to compare?"))
    
    papers=[]
    
    for i in range(num_papers):
    
        pdf_path=input(f"Enter PDF path for paper {i+1}: ")
        article=extract_pdf_text(pdf_path)
        papers.append(article)
    
    summaries=[]
    
    for i,paper in enumerate(papers):
    
        print(f"Summarizing paper{i+1}/{len(papers)}")
    
        prompt= f"""
        Summarize this research paper.

        Include:
        - Objective
        - Methodology
        - Key Findings
        - Limitations

        Paper:
        {paper}
        """
        summary=generate_response(prompt)
        summaries.append(summary)
    
    combined_summaries=""
    
    for i, summary in enumerate(summaries):
        combined_summaries+=f"\nPaper{i+1}\n"
        combined_summaries+=summary
    
    comparison_prompt = f"""
    Compare all these research papers.

    For each paper discuss:

    - Objective
    - Methodology
    - Strengths
    - Weaknesses

    Then provide:

    1. Similarities
    2. Differences
    3. Research Gaps
    4. Future Scope

    Papers:

    {combined_summaries}
    """

    output=generate_response(comparison_prompt)
    
    print(output)
    
    exit()


elif choice=="9":
    num_papers=int(input("How many papers do you want to compare?"))
    papers=[]
    paper_summaries=[]

    for i in range(num_papers):
    
        pdf_path=input(f"Enter PDF path for paper {i+1}: ")
        article=extract_pdf_text(pdf_path)
        article = article[:15000]
        papers.append(article)
        summary_prompt = f"""
        Summarize this research paper.

        Include:
        - Objective
        - Methodology
        - Key Findings
        - Limitations

        Paper:
        {article}
        """
        summary=generate_response(summary_prompt)
        paper_summaries.append(summary)
    
    all_papers="\n\n".join(paper_summaries)

    prompt = f"""
    Generate a structured literature review from these research papers.

    Include:

    1. Introduction
    2. Existing Research
    3. Methodologies Used
    4. Comparative Analysis
    5. Research Gaps
    6. Future Scope
    7. Conclusion

    Research Papers:

    {all_papers}
    """
    output=generate_response(prompt)
    print(output)
    exit()


elif choice=="10":
    prompt=research_gap_prompt(article)


elif choice=="11":
    prompt = recommendation_prompt(article)

elif choice=="12":
    prompt=citation_prompt(article)

elif choice=="13":
    prompt=metadata_prompt(article)

else :
    print("invalid input")
    exit()

if len(article) < 8000:

    print("Small paper detected")

    final_output = generate_response(prompt)

else:

    chunks = chunk_text(article)
    index, chunks = create_vector_store(chunks)

    task_prompt = prompt

    all_summaries = ""

    print(f"\nTotal Chunks: {len(chunks)}")

    for index, chunk in enumerate(chunks):

        current_prompt = task_prompt.replace(article, chunk)

        print(f"\nProcessing Chunk {index+1}/{len(chunks)}")

        summary = generate_response(current_prompt)

        all_summaries += summary + "\n"

    print("\nGenerating Final Summary...")

    final_prompt = f"""
    Combine and refine the following outputs into one coherent response.

    Outputs:
    {all_summaries}
    """

    final_output = generate_response(final_prompt)

print("\n=====FINAL RESPONSE=====\n")


print(final_output)

filename="output.txt"

if choice=="10":
    filename="research_gap_analysis.txt"

elif choice == "11":
    filename = "recommendations.txt"

elif choice == "12":
    filename = "citations.txt"

with open(filename, "w", encoding="utf-8") as save:
    save.write(final_output)

doc=Document()

doc.add_heading("Research Paper Analysis", level=1)

for line in final_output.split("\n"):
    if line.strip():
        doc.add_paragraph(line)

doc_filename=filename.replace(".txt",".docx")

doc.save(doc_filename)

print(f"TXT saved as {filename}")
print(f"DOCX saved as{doc_filename}")

