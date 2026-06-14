from docx import Document

def save_docx(filename , content):
    doc=Document()
    doc.add_heading("Research Paper Analysis", level=1)
    doc.add_paragraph(content)
    doc.save(filename)
    